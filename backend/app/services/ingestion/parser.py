import os
import re
import uuid
from abc import ABC, abstractmethod
from collections import defaultdict
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from app.core.logging import logger
from app.domain.parser import (
    ParsedBlock,
    ParsedBlockType,
    ParsedDocument,
    ParsedPage,
    ParsedTable,
)


class BaseDocumentParser(ABC):
    """Abstract interface for all document parsing engines in Selnikel AI."""

    @abstractmethod
    async def parse(self, file_path: str, content_type: Optional[str] = None) -> ParsedDocument:
        """Parse a document file into structured ParsedDocument with page-level attribution."""
        pass

    @abstractmethod
    def supports(self, file_path: str, content_type: Optional[str] = None) -> bool:
        """Check if this parser supports the given file format."""
        pass


class FastFallbackParser(BaseDocumentParser):
    """Lightweight, resilient fallback parser handling plain text, markdown, CSV, DOCX, and PDFs.
    Guarantees structured table extraction, section hierarchy, and page provenance.
    """

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".log", ".pdf", ".docx"}

    def supports(self, file_path: str, content_type: Optional[str] = None) -> bool:
        ext = Path(file_path).suffix.lower()
        if ext in self.SUPPORTED_EXTENSIONS:
            return True
        if content_type and any(
            t in content_type.lower()
            for t in ["text/plain", "text/markdown", "text/csv", "application/pdf", "wordprocessingml"]
        ):
            return True
        return False

    def parse_sync(self, file_path: str, content_type: Optional[str] = None) -> ParsedDocument:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document file not found: {file_path}")

        filename = path.name
        ext = path.suffix.lower()

        if ext == ".pdf" or (content_type and "pdf" in content_type.lower()):
            return self._parse_pdf(path)
        elif ext == ".docx" or (content_type and "word" in content_type.lower()):
            return self._parse_docx(path)
        else:
            return self._parse_text(path)

    async def parse(self, file_path: str, content_type: Optional[str] = None) -> ParsedDocument:
        return self.parse_sync(file_path, content_type=content_type)

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        path = Path(path)
        try:
            import pypdf

            reader = pypdf.PdfReader(str(path))
            total_pages = len(reader.pages)
            pages: List[ParsedPage] = []
            blocks: List[ParsedBlock] = []
            all_text_parts: List[str] = []
            all_tables: List[ParsedTable] = []
            all_rejected_rows: List[Dict[str, Any]] = []

            table_global_idx = 1
            for idx, page in enumerate(reader.pages, start=1):
                page_tables, page_text, section_headers, rejected_rows = self._extract_pdf_page_content(
                    page, page_number=idx, start_table_idx=table_global_idx
                )
                table_global_idx += len(page_tables)
                all_tables.extend(page_tables)
                all_rejected_rows.extend(rejected_rows)

                all_text_parts.append(f"<!-- Page {idx} -->\n{page_text}")

                pages.append(
                    ParsedPage(
                        page_number=idx,
                        text_content=page_text,
                        tables=page_tables,
                        section_headers=section_headers,
                    )
                )

                if page_text:
                    blocks.append(
                        ParsedBlock(
                            content=page_text,
                            block_type=ParsedBlockType.PARAGRAPH,
                            page_number=idx,
                        )
                    )

            if all_rejected_rows:
                safe_telemetry = [
                    {
                        "page_number": r["page_number"],
                        "reason": r["reason"],
                        "expected_cols": r["expected_cols"],
                        "actual_cols": r["actual_cols"],
                    }
                    for r in all_rejected_rows
                ]
                logger.warning(
                    f"FastFallbackParser detected {len(all_rejected_rows)} non-conforming table rows across {path.name}: {safe_telemetry}"
                )

            full_markdown = "\n\n".join(all_text_parts)
            return ParsedDocument(
                filename=path.name,
                total_pages=max(1, total_pages),
                full_markdown=full_markdown,
                pages=pages,
                tables=all_tables,
                blocks=blocks,
                metadata={
                    "file_size_bytes": path.stat().st_size,
                    "ocr_applied": False,
                    "parser_name": "fast_fallback_pypdf",
                    "rejected_table_rows": all_rejected_rows,
                    "table_parsing_warnings": [
                        f"Page {r['page_number']}: {r['reason']} (expected {r['expected_cols']}, got {r['actual_cols']})"
                        for r in all_rejected_rows
                    ],
                },
                parser_name="fast_fallback_pypdf",
            )
        except Exception as e:
            logger.error(f"FastFallbackParser PDF extraction failed: {e}")
            raise

    def _extract_pdf_page_content(
        self, page, page_number: int, start_table_idx: int = 1
    ) -> Tuple[List[ParsedTable], str, List[str], List[Dict[str, Any]]]:
        """Extracts spatial text elements, detects sections and tabular layouts from PDF pages with zero silent loss."""
        raw_text = page.extract_text() or ""
        section_headers: List[str] = []
        rejected_rows: List[Dict[str, Any]] = []

        # 1. Extract Section Headers reliably from page text (excluding model codes and table values)
        for line in raw_text.splitlines():
            line_str = line.strip()
            if not line_str:
                continue
            if re.match(r"^\d+(\.\d+)*\.?\s+[A-Z]", line_str):
                if line_str not in section_headers:
                    section_headers.append(line_str)
            elif (
                line_str.isupper()
                and 5 <= len(line_str) <= 60
                and not re.search(r"\d", line_str)
                and "SELNIKEL" not in line_str
                and not line_str.startswith("SB-")
            ):
                if line_str not in section_headers:
                    section_headers.append(line_str)

        # 2. Extract spatial elements for table reconstruction
        elements = []

        def visitor(text, cm, tm, font_dict, font_size):
            cleaned = text.strip()
            if cleaned and tm is not None and len(tm) >= 6:
                x = tm[4]
                y = tm[5]
                elements.append((y, x, cleaned))

        try:
            page.extract_text(visitor_text=visitor)
        except Exception:
            elements = []

        if not elements:
            return [], raw_text.strip(), section_headers, []

        # Group elements by y-coordinate within vertical tolerance (3.0 points)
        lines_dict = defaultdict(list)
        for y, x, text in elements:
            matched_y = None
            for existing_y in lines_dict:
                if abs(existing_y - y) <= 3.0:
                    matched_y = existing_y
                    break
            if matched_y is None:
                matched_y = y
            lines_dict[matched_y].append((x, text))

        sorted_y = sorted(lines_dict.keys(), reverse=True)
        multi_col_rows = []

        for y in sorted_y:
            line_elems = sorted(lines_dict[y], key=lambda item: item[0])
            line_texts = [t for _, t in line_elems]
            full_line = " ".join(line_texts).strip()

            # Skip full-line section headers or title banners from table candidates
            is_section_line = any(full_line == sh.strip() for sh in section_headers) or re.match(
                r"^\d+(\.\d+)*\.?\s+[A-Z]", full_line
            )
            if not is_section_line and len(line_texts) >= 3:
                multi_col_rows.append((y, line_texts))

        # Spatial table clustering: group multi-column rows by vertical proximity (dy <= 35 pt)
        # Robust against intervening mismatched rows: does not shred subsequent valid rows
        table_clusters: List[List[Tuple[float, List[str]]]] = []
        curr_cluster: List[Tuple[float, List[str]]] = []
        last_y: Optional[float] = None

        for y, row in multi_col_rows:
            if not curr_cluster:
                curr_cluster.append((y, row))
                last_y = y
            else:
                hdr_col_count = len(curr_cluster[0][1])
                dy = abs(last_y - y) if last_y is not None else 0.0

                if dy > 35.0:
                    if len(curr_cluster) >= 2:
                        table_clusters.append(curr_cluster)
                    elif len(curr_cluster) == 1:
                        rejected_rows.append({
                            "page_number": page_number,
                            "y": curr_cluster[0][0],
                            "row_cells": curr_cluster[0][1],
                            "expected_cols": len(curr_cluster[0][1]),
                            "actual_cols": len(curr_cluster[0][1]),
                            "reason": "orphan_single_row_no_table_body",
                        })
                    curr_cluster = [(y, row)]
                    last_y = y
                else:
                    if len(row) == hdr_col_count:
                        curr_cluster.append((y, row))
                        last_y = y
                    else:
                        # Record rejected row with structured diagnostic and continue active cluster
                        rejected_rows.append({
                            "page_number": page_number,
                            "y": y,
                            "row_cells": row,
                            "expected_cols": hdr_col_count,
                            "actual_cols": len(row),
                            "reason": "column_count_mismatch",
                        })
                        last_y = y

        if len(curr_cluster) >= 2:
            table_clusters.append(curr_cluster)
        elif len(curr_cluster) == 1:
            rejected_rows.append({
                "page_number": page_number,
                "y": curr_cluster[0][0],
                "row_cells": curr_cluster[0][1],
                "expected_cols": len(curr_cluster[0][1]),
                "actual_cols": len(curr_cluster[0][1]),
                "reason": "orphan_single_row_no_table_body",
            })

        # Reconstruct clean GFM tables from clusters
        tables: List[ParsedTable] = []
        for cluster_idx, cluster in enumerate(table_clusters, start=1):
            table_rows = [r for _, r in cluster]
            hdr = table_rows[0]
            col_count = len(hdr)
            valid_body_rows = table_rows[1:]

            if valid_body_rows and col_count >= 2:
                sep = ["---"] * col_count
                header_line = "| " + " | ".join(hdr) + " |"
                sep_line = "| " + " | ".join(sep) + " |"
                body_lines = ["| " + " | ".join(r) + " |" for r in valid_body_rows]

                table_md = "\n".join([header_line, sep_line] + body_lines)
                tables.append(
                    ParsedTable(
                        table_id=f"pdf_tab_{start_table_idx + cluster_idx - 1:02d}",
                        page_number=page_number,
                        markdown_table=table_md,
                        num_rows=len(valid_body_rows),
                        num_cols=col_count,
                        headers=hdr,
                        caption=section_headers[-1] if section_headers else f"Table Page {page_number}",
                    )
                )

        page_text = raw_text.strip()
        return tables, page_text, section_headers, rejected_rows

    def _parse_docx(self, path: Path) -> ParsedDocument:
        """Parses DOCX preserving document flow order, page breaks, and tabular layout."""
        try:
            import docx
            path = Path(path)
            doc = docx.Document(str(path))
            pages: List[ParsedPage] = []
            blocks: List[ParsedBlock] = []
            all_tables: List[ParsedTable] = []

            curr_page_num = 1
            curr_page_text_lines: List[str] = []
            curr_page_tables: List[ParsedTable] = []
            curr_page_headers: List[str] = []

            def finalize_page():
                nonlocal curr_page_num, curr_page_text_lines, curr_page_tables, curr_page_headers
                if curr_page_text_lines or curr_page_tables:
                    page_full_text = "\n\n".join(curr_page_text_lines)
                    pages.append(
                        ParsedPage(
                            page_number=curr_page_num,
                            text_content=page_full_text,
                            tables=curr_page_tables,
                            section_headers=curr_page_headers,
                        )
                    )
                    curr_page_num += 1
                    curr_page_text_lines = []
                    curr_page_tables = []
                    curr_page_headers = []

            W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            for child in doc.element.body:
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

                if tag == "p":
                    p = docx.text.paragraph.Paragraph(child, doc)

                    # 1. Check pageBreakBefore (must finalize page BEFORE adding paragraph content)
                    has_page_break_before = False
                    try:
                        if p.paragraph_format.page_break_before:
                            has_page_break_before = True
                    except Exception:
                        pass
                    if not has_page_break_before:
                        for pPr in child.iter(f"{{{W_NS}}}pPr"):
                            if any(c.tag.endswith("pageBreakBefore") for c in pPr):
                                has_page_break_before = True
                                break

                    if has_page_break_before:
                        finalize_page()

                    # 2. Extract segments separated by inline page breaks
                    segments: List[str] = []
                    current_parts: List[str] = []

                    for elem in child.iter():
                        elem_tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
                        if elem_tag == "br":
                            br_type = elem.get(f"{{{W_NS}}}type") or elem.get("type") or elem.attrib.get("w:type")
                            if br_type == "page":
                                seg = "".join(current_parts).strip()
                                segments.append(seg)
                                current_parts = []
                            else:
                                current_parts.append("\n")
                        elif elem_tag == "cr":
                            current_parts.append("\n")
                        elif elem_tag == "t" and elem.text:
                            current_parts.append(elem.text)
                        elif elem_tag == "tab":
                            current_parts.append("\t")

                    final_seg = "".join(current_parts).strip()
                    segments.append(final_seg)

                    is_heading = (p.style and p.style.name.startswith("Heading")) or bool(
                        re.match(r"^\d+(\.\d+)*\.?\s+[A-Z]", p.text.strip())
                    )

                    for i, seg in enumerate(segments):
                        if i > 0:
                            # Crossed inline page break!
                            finalize_page()
                        if seg:
                            curr_page_text_lines.append(seg)
                            if is_heading and seg not in curr_page_headers:
                                curr_page_headers.append(seg)
                            blocks.append(
                                ParsedBlock(
                                    content=seg,
                                    block_type=ParsedBlockType.HEADING if is_heading else ParsedBlockType.PARAGRAPH,
                                    page_number=curr_page_num,
                                )
                            )

                elif tag == "tbl":
                    tbl = docx.table.Table(child, doc)
                    rows_data = []
                    for row in tbl.rows:
                        # Escape pipe characters inside cell text to preserve GFM table syntax
                        row_cells = [cell.text.strip().replace("\n", " ").replace("|", "\\|") for cell in row.cells]
                        rows_data.append(row_cells)

                    if rows_data and len(rows_data) >= 2:
                        headers = rows_data[0]
                        col_count = len(headers)
                        header_line = "| " + " | ".join(headers) + " |"
                        sep_line = "| " + " | ".join(["---"] * col_count) + " |"
                        body_lines = ["| " + " | ".join(r) + " |" for r in rows_data[1:]]
                        table_md = "\n".join([header_line, sep_line] + body_lines)

                        parsed_tab = ParsedTable(
                            table_id=f"docx_tab_{len(all_tables)+1:02d}",
                            page_number=curr_page_num,
                            markdown_table=table_md,
                            num_rows=len(rows_data) - 1,
                            num_cols=col_count,
                            headers=headers,
                            caption=curr_page_headers[-1] if curr_page_headers else f"Table Page {curr_page_num}",
                        )
                        curr_page_tables.append(parsed_tab)
                        all_tables.append(parsed_tab)
                        curr_page_text_lines.append(table_md)
                        blocks.append(
                            ParsedBlock(
                                content=table_md,
                                block_type=ParsedBlockType.TABLE,
                                page_number=curr_page_num,
                            )
                        )

            # Finalize remaining page
            finalize_page()

            if not pages:
                pages.append(
                    ParsedPage(
                        page_number=1,
                        text_content="",
                        tables=[],
                        section_headers=[],
                    )
                )

            full_markdown = "\n\n".join(p.text_content for p in pages)
            return ParsedDocument(
                filename=path.name,
                total_pages=len(pages),
                full_markdown=full_markdown,
                pages=pages,
                tables=all_tables,
                blocks=blocks,
                metadata={
                    "file_size_bytes": path.stat().st_size,
                    "ocr_applied": False,
                    "parser_name": "fast_fallback_docx",
                },
                parser_name="fast_fallback_docx",
            )
        except Exception as e:
            logger.error(f"FastFallbackParser DOCX extraction failed: {e}")
            raise

    def _parse_text(self, path: Path) -> ParsedDocument:
        path = Path(path)
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        tables = self._extract_markdown_tables(content, page_number=1)
        total_pages = 1

        section_headers = [
            line.strip()
            for line in content.splitlines()
            if line.strip().startswith("#") or re.match(r"^\d+(\.\d+)*\.?\s+[A-Z]", line.strip())
        ]

        page = ParsedPage(
            page_number=1,
            text_content=content,
            tables=tables,
            section_headers=section_headers,
        )

        return ParsedDocument(
            filename=path.name,
            total_pages=total_pages,
            full_markdown=content,
            pages=[page],
            tables=tables,
            blocks=[
                ParsedBlock(
                    content=content,
                    block_type=ParsedBlockType.PARAGRAPH,
                    page_number=1,
                )
            ],
            metadata={
                "file_size_bytes": path.stat().st_size,
                "ocr_applied": False,
                "parser_name": "fast_fallback_text",
            },
            parser_name="fast_fallback_text",
        )

    def _extract_markdown_tables(self, content: str, page_number: int = 1) -> List[ParsedTable]:
        tables: List[ParsedTable] = []
        table_pattern = re.compile(
            r"(\|[^\n]+\|\r?\n\|[-:\s|]+\|\r?\n(?:\|[^\n]+\|\r?\n?)+)", re.MULTILINE
        )
        for match in table_pattern.finditer(content):
            table_md = match.group(0).strip()
            lines = [l.strip() for l in table_md.splitlines() if l.strip()]
            if len(lines) >= 3:
                headers = [c.strip() for c in lines[0].split("|")[1:-1]]
                preceding_text = content[:match.start()].strip()
                preceding_lines = [l.strip() for l in preceding_text.splitlines() if l.strip()]
                caption = None
                if preceding_lines:
                    last_line = preceding_lines[-1]
                    caption = last_line.lstrip("#").strip()

                tables.append(
                    ParsedTable(
                        table_id=f"txt_tab_{len(tables)+1:02d}",
                        page_number=page_number,
                        markdown_table=table_md,
                        num_rows=len(lines) - 2,
                        num_cols=len(headers),
                        headers=headers,
                        caption=caption or f"Table Page {page_number}",
                    )
                )
        return tables


class DoclingParser(BaseDocumentParser):
    """Advanced industrial document parser using IBM Docling for structural layout
    and table extraction. Preserves technical tables in GitHub-Flavored Markdown.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".pptx", ".html", ".asciidoc", ".md"}

    def __init__(self):
        self._converter = None
        self._docling_available = False
        self._init_docling()

    def _init_docling(self) -> None:
        try:
            from docling.document_converter import DocumentConverter

            self._converter = DocumentConverter()
            self._docling_available = True
            logger.info("Docling DocumentConverter initialized successfully.")
        except ImportError:
            self._docling_available = False
            logger.warning(
                "Docling package not installed. Parser will automatically route to FastFallbackParser."
            )
        except Exception as e:
            self._docling_available = False
            logger.warning(f"Docling initialization error: {e}. Fallback enabled.")

    @property
    def is_available(self) -> bool:
        return self._docling_available

    def supports(self, file_path: str, content_type: Optional[str] = None) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS

    async def parse(self, file_path: str, content_type: Optional[str] = None) -> ParsedDocument:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document file not found: {file_path}")

        # If Docling is not available, delegate seamlessly to fallback
        if not self._docling_available:
            logger.info(f"Docling unavailable, delegating {path.name} to FastFallbackParser")
            fallback = FastFallbackParser()
            return await fallback.parse(file_path, content_type)

        try:
            conv_result = self._converter.convert(str(path))
            doc = conv_result.document

            full_markdown = doc.export_to_markdown()
            tables: List[ParsedTable] = []
            blocks: List[ParsedBlock] = []

            # Extract tables directly from Docling structure
            if hasattr(doc, "tables") and doc.tables:
                for t_idx, table_item in enumerate(doc.tables):
                    page_no = 1
                    if hasattr(table_item, "prov") and table_item.prov:
                        page_no = getattr(table_item.prov[0], "page_no", 1)

                    table_md = table_item.export_to_markdown() if hasattr(table_item, "export_to_markdown") else str(table_item)
                    tables.append(
                        ParsedTable(
                            table_id=str(uuid.uuid4()),
                            page_number=page_no,
                            markdown_table=table_md,
                            caption=getattr(table_item, "caption", None),
                            num_rows=getattr(table_item, "num_rows", 0),
                            num_cols=getattr(table_item, "num_cols", 0),
                        )
                    )
                    blocks.append(
                        ParsedBlock(
                            content=table_md,
                            block_type=ParsedBlockType.TABLE,
                            page_number=page_no,
                        )
                    )

            total_pages = 1
            if hasattr(doc, "pages") and doc.pages:
                total_pages = len(doc.pages)

            page = ParsedPage(
                page_number=1,
                text_content=full_markdown,
                tables=tables,
                section_headers=[],
            )

            return ParsedDocument(
                filename=path.name,
                total_pages=total_pages,
                full_markdown=full_markdown,
                pages=[page],
                tables=tables,
                blocks=blocks,
                metadata={
                    "file_size_bytes": path.stat().st_size,
                    "ocr_applied": False,
                    "parser_name": "docling",
                },
                parser_name="docling",
            )
        except Exception as e:
            logger.error(f"Docling parsing failed for {path.name}: {e}. Routing to fallback parser.")
            fallback = FastFallbackParser()
            return await fallback.parse(file_path, content_type)


class DocumentParserFactory:
    """Factory selecting the optimal parser (Docling vs FastFallback) based on file type and availability."""

    @staticmethod
    def get_parser(
        file_path: str,
        content_type: Optional[str] = None,
        force_fallback: bool = False,
    ) -> BaseDocumentParser:
        if force_fallback:
            return FastFallbackParser()

        docling = DoclingParser()
        if docling.is_available and docling.supports(file_path, content_type):
            return docling

        return FastFallbackParser()


document_parser_factory = DocumentParserFactory()
