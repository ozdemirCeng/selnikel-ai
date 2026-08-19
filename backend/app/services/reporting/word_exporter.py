import io
import re
from typing import List
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


class EngineeringWordExporter:
    """Exports markdown engineering reports to formal Microsoft Word (.docx) documents."""

    @classmethod
    def generate_docx(cls, markdown_text: str, title: str = "Selnikel Teknik Raporu") -> bytes:
        doc = docx.Document()

        # Page Margins
        sections = doc.sections
        for s in sections:
            s.top_margin = Inches(0.8)
            s.bottom_margin = Inches(0.8)
            s.left_margin = Inches(0.8)
            s.right_margin = Inches(0.8)

        # Header Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"SELNİKEL ENERJİ — {title.upper()}")
        title_run.font.name = "Arial"
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(15, 23, 42)

        sub_para = doc.add_paragraph()
        sub_run = sub_para.add_run("AR-GE ve Mühendislik Bilgi Sistemi &bull; Resmi Teknik Dokümantasyon")
        sub_run.font.name = "Arial"
        sub_run.font.size = Pt(9)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(2, 132, 199)

        doc.add_paragraph("―" * 45)

        lines = markdown_text.split("\n")
        table_rows: List[List[str]] = []
        in_table = False

        for line in lines:
            trimmed = line.strip()

            if trimmed.startswith("|") and trimmed.endswith("|"):
                if re.match(r"^\|[\s\-:|]+\|$", trimmed):
                    continue
                cols = [c.strip() for c in trimmed.strip("|").split("|")]
                table_rows.append(cols)
                in_table = True
                continue
            elif in_table and table_rows:
                cls._render_word_table(doc, table_rows)
                table_rows = []
                in_table = False

            if not trimmed:
                continue

            if trimmed.startswith("# "):
                h = doc.add_heading(trimmed[2:], level=1)
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(4)
            elif trimmed.startswith("## "):
                h = doc.add_heading(trimmed[3:], level=2)
                h.paragraph_format.space_before = Pt(10)
                h.paragraph_format.space_after = Pt(3)
            elif trimmed.startswith("### "):
                h = doc.add_heading(trimmed[4:], level=3)
                h.paragraph_format.space_before = Pt(8)
                h.paragraph_format.space_after = Pt(2)
            elif trimmed.startswith("- ") or trimmed.startswith("* "):
                p = doc.add_paragraph(trimmed[2:], style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
            elif trimmed == "---":
                doc.add_paragraph("―" * 45)
            else:
                p = doc.add_paragraph(trimmed)
                p.paragraph_format.space_after = Pt(4)

        if table_rows:
            cls._render_word_table(doc, table_rows)

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        return docx_bytes

    @classmethod
    def _render_word_table(cls, doc: docx.Document, rows: List[List[str]]) -> None:
        if not rows:
            return

        num_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx]
            is_header = r_idx == 0
            for c_idx, val in enumerate(row_data):
                if c_idx < len(row.cells):
                    cell = row.cells[c_idx]
                    cell.text = val
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(9.5)
                            if is_header:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)

                    if is_header:
                        tcPr = cell._tc.get_or_add_tcPr()
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:val"), "clear")
                        shd.set(qn("w:color"), "auto")
                        shd.set(qn("w:fill"), "0369A1")
                        tcPr.append(shd)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)
