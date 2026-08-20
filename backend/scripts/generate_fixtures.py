import os
import hashlib
import json
from pathlib import Path
import docx
from docx.enum.text import WD_BREAK
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

DOCS_DIR = Path('backend/tests/fixtures/documents')
DOCS_DIR.mkdir(parents=True, exist_ok=True)
MANIFEST_PATH = Path('backend/tests/fixtures/fixture_manifest.json')

def generate_pdf_rev01():
    pdf_path = DOCS_DIR / 'SB_Series_Steam_Boiler_Datasheet_REV01.pdf'
    doc = SimpleDocTemplate(str(pdf_path), pagesize=letter, leftMargin=36, rightMargin=36, topMargin=54, bottomMargin=54)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('DocTitle', parent=styles['Heading1'], fontSize=15, leading=18, alignment=1)
    sec_style = ParagraphStyle('SecHdr', parent=styles['Heading2'], fontSize=12, leading=15, spaceBefore=6, spaceAfter=12)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9, leading=12)

    story = []
    # Page 1
    story.append(Paragraph('SELNIKEL ISITMA SISTEMLERI A.S.', title_style))
    story.append(Paragraph('SB-Series Industrial Steam Boiler Technical Datasheet (REV-01)', title_style))
    story.append(Paragraph('Document Code: SB-DS-2025 | Revision: REV-01 | Department: Thermal Engineering', body_style))
    story.append(Spacer(1, 15))
    story.append(Paragraph('1. General Overview & Design Standards', sec_style))
    story.append(Paragraph('The SB-Series steam boiler baseline design operates with dual-pass Scotch marine shell geometry.', body_style))
    story.append(PageBreak())

    # Page 2
    story.append(Paragraph('2. Technical Operating Parameters & Capacities', sec_style))
    story.append(Spacer(1, 10))
    table_data = [
        ['Model Code', 'Steam Cap. (t/h)', 'Working Press. (bar)', 'Design Press. (bar)', 'Steam Temp. (°C)', 'Thermal Power (kW)'],
        ['SB-500', '0.45 t/h', '10.0 bar', '14.0 bar', '184.1 °C', '315 kW'],
        ['SB-1000', '0.9 t/h', '14.0 bar', '16.0 bar', '198.3 °C', '630 kW'],
        ['SB-2000', '1.8 t/h', '14.0 bar', '18.0 bar', '198.3 °C', '1260 kW'],
        ['SB-5000', '4.5 t/h', '14.0 bar', '22.0 bar', '198.3 °C', '3150 kW'],
    ]
    t = Table(table_data, colWidths=[65, 85, 95, 95, 85, 95])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#2C3E50')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 8),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
    ]))
    story.append(t)
    story.append(PageBreak())

    # Page 3
    story.append(Paragraph('3. Safety Relief Valves & Auxiliary Limits', sec_style))
    story.append(Spacer(1, 10))
    v_data = [
        ['Safety Device', 'Set Pressure (bar)', 'Connection DN (mm)', 'Discharge Cap. (kg/h)'],
        ['Primary Safety Valve', '14.5 bar', '32 mm', '1100 kg/h'],
        ['Secondary Safety Valve', '15.0 bar', '40 mm', '1350 kg/h'],
    ]
    vt = Table(v_data, colWidths=[130, 100, 110, 120])
    vt.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#2C3E50')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 8),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
    ]))
    story.append(vt)
    doc.build(story)
    print(f'Generated: {pdf_path}')

def generate_pdf_rev02():
    pdf_path = DOCS_DIR / 'SB_Series_Steam_Boiler_Datasheet.pdf'
    doc = SimpleDocTemplate(str(pdf_path), pagesize=letter, leftMargin=36, rightMargin=36, topMargin=54, bottomMargin=54)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('DocTitle', parent=styles['Heading1'], fontSize=15, leading=18, alignment=1)
    sec_style = ParagraphStyle('SecHdr', parent=styles['Heading2'], fontSize=12, leading=15, spaceBefore=6, spaceAfter=12)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9, leading=12)

    story = []
    story.append(Paragraph('SELNIKEL ISITMA SISTEMLERI A.S.', title_style))
    story.append(Paragraph('SB-Series Industrial Steam Boiler Technical Datasheet', title_style))
    story.append(Paragraph('Document Code: SB-DS-2026 | Revision: REV-02 | Department: Thermal Engineering', body_style))
    story.append(Spacer(1, 15))
    story.append(Paragraph('1. General Overview & Design Standards', sec_style))
    story.append(Paragraph('The SB-Series steam boiler represents three-pass wet-back Scotch marine construction.', body_style))
    story.append(PageBreak())

    story.append(Paragraph('2. Technical Operating Parameters & Capacities', sec_style))
    story.append(Spacer(1, 10))
    table_data = [
        ['Model Code', 'Steam Cap. (t/h)', 'Working Press. (bar)', 'Design Press. (bar)', 'Steam Temp. (°C)', 'Thermal Power (kW)'],
        ['SB-500', '0.5 t/h', '12.0 bar', '16.0 bar', '191.6 °C', '350 kW'],
        ['SB-1000', '1.0 t/h', '16.0 bar', '18.0 bar', '204.3 °C', '700 kW'],
        ['SB-2000', '2.0 t/h', '16.0 bar', '20.0 bar', '204.3 °C', '1400 kW'],
        ['SB-5000', '5.0 t/h', '16.0 bar', '25.0 bar', '204.3 °C', '3500 kW'],
    ]
    t = Table(table_data, colWidths=[65, 85, 95, 95, 85, 95])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#2C3E50')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 8),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
    ]))
    story.append(t)
    story.append(PageBreak())

    story.append(Paragraph('3. Safety Relief Valves & Auxiliary Limits', sec_style))
    story.append(Spacer(1, 10))
    v_data = [
        ['Safety Device', 'Set Pressure (bar)', 'Connection DN (mm)', 'Discharge Cap. (kg/h)'],
        ['Primary Safety Valve', '16.5 bar', '32 mm', '1250 kg/h'],
        ['Secondary Safety Valve', '17.0 bar', '40 mm', '1500 kg/h'],
    ]
    vt = Table(v_data, colWidths=[130, 100, 110, 120])
    vt.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#2C3E50')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 8),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
    ]))
    story.append(vt)
    doc.build(story)
    print(f'Generated: {pdf_path}')

def generate_burner_docx():
    docx_path = DOCS_DIR / 'Monoblock_Burner_Maintenance_Manual.docx'
    doc = docx.Document()
    
    # Page 1
    doc.add_heading('SELNIKEL MONOBLOCK INDUSTRIAL GAS BURNER', level=1)
    doc.add_paragraph('Technical Maintenance and Fault Diagnostic Manual | Revision: REV-01 | Model: MB-Series')
    
    doc.add_heading('1. Technical Air & Combustion Data', level=2)
    doc.add_paragraph('Monoblock industrial gas burners are engineered for high-efficiency thermal plants with electronic ratio control.')
    
    doc.add_heading('2. Periodic Maintenance Schedule & Intervals', level=2)
    tbl = doc.add_table(rows=5, cols=4)
    tbl.style = 'Table Grid'
    headers = ['Component', 'Inspection Action', 'Service Interval', 'Torque / Spec']
    for i, h in enumerate(headers):
        tbl.cell(0, i).text = h
    
    rows = [
        ['Burner Nozzles', 'Clean and inspect spray pattern', '500 hour', '25 Nm'],
        ['Nozzle Overhaul', 'Replace atomizing disc and filter', '2000 hour', '30 Nm'],
        ['Gas Valve Train', 'Tightness and leak test', '6 month', 'No bubble at 100 mbar'],
        ['Complete Blower Overhaul', 'Bearing lubrication and dynamic balancing', '1 year', 'ISO 1940 G2.5'],
    ]
    for row_idx, r in enumerate(rows, start=1):
        for col_idx, val in enumerate(r):
            tbl.cell(row_idx, col_idx).text = val
            
    # Page Break to Page 2
    p_break = doc.add_paragraph()
    p_break.add_run().add_break(WD_BREAK.PAGE)
    
    # Page 2
    doc.add_heading('3. Diagnostic Fault Codes & Troubleshooting Matrix', level=2)
    doc.add_paragraph('When a burner safety lock occurs, consult the diagnostic fault code displayed on the digital control unit:')
    
    tbl_faults = doc.add_table(rows=4, cols=4)
    tbl_faults.style = 'Table Grid'
    fault_headers = ['Fault Code', 'Fault Description', 'Possible Root Cause', 'Action Required']
    for i, h in enumerate(fault_headers):
        tbl_faults.cell(0, i).text = h
    
    fault_rows = [
        ['E01', 'Flame Failure Lockout', 'Ionization probe contamination or gas cutoff', 'Clean ionization probe and verify supply pressure'],
        ['E04', 'Air Pressure Switch Fault', 'Insufficient combustion air flow or clogged fan', 'Check blower differential pressure and air switch tube'],
        ['E07', 'High Flue Gas Temperature Trip', 'Heat exchanger soot buildup or excess firing rate', 'Shutdown burner and perform boiler tube cleaning'],
    ]
    for row_idx, r in enumerate(fault_rows, start=1):
        for col_idx, val in enumerate(r):
            tbl_faults.cell(row_idx, col_idx).text = val
            
    doc.add_heading('4. Burner Nozzle & Mechanical Parts Compatibility', level=2)
    doc.add_paragraph('Mechanical connection standards for MB-Series monoblock burners:')
    doc.add_paragraph('Burner Nozzle Thread Specification: 9/16-24 UNEF standard thread.')
    doc.add_paragraph('Oil and Gas Filter Element Mesh: 100 microns rated stainless steel wire.')
    doc.add_paragraph('Flange Connection Interface: DN65 PN16 standard industrial gas inlet.')
    
    doc.save(str(docx_path))
    print(f'Generated: {docx_path}')

def generate_commissioning_docx():
    docx_path = DOCS_DIR / 'Industrial_Boiler_Commissioning_Guide.docx'
    doc = docx.Document()
    
    # Page 1
    doc.add_heading('SELNIKEL INDUSTRIAL BOILER COMMISSIONING GUIDE', level=1)
    doc.add_paragraph('Site Installation, Electrical Verification & Regulatory Standards | Revision: REV-01')
    
    doc.add_heading('1. Pre-Commissioning Electrical Limits', level=2)
    tbl_elec = doc.add_table(rows=3, cols=3)
    tbl_elec.style = 'Table Grid'
    headers_elec = ['Terminal', 'Signal Type', 'Allowed Range']
    for i, h in enumerate(headers_elec):
        tbl_elec.cell(0, i).text = h
    elec_rows = [
        ['L1-L2-L3', 'Main Power', '380 - 420 V'],
        ['P1-P2', 'Pressure 4-20 mA', '0.0 - 25.0 bar'],
    ]
    for row_idx, r in enumerate(elec_rows, start=1):
        for col_idx, val in enumerate(r):
            tbl_elec.cell(row_idx, col_idx).text = val
            
    doc.add_heading('2. Standards, Directives & Regulatory Compliance', level=2)
    doc.add_paragraph('All Selnikel industrial steam boilers and burner installations strictly comply with mandatory European harmonized standards:')
    doc.add_paragraph('Shell Boilers Manufacturing Standard: EN 12953 for shell boilers design, safety and manufacturing requirements.')
    doc.add_paragraph('Pressure Equipment Directive: PED 2014/68/EU Category IV conformity module H1 certification.')
    doc.add_paragraph('Forced Draught Gas Burner Standard: EN 676 automatic forced draught burners for gaseous fuels.')
    
    # Page Break to Page 2
    p_break = doc.add_paragraph()
    p_break.add_run().add_break(WD_BREAK.PAGE)
    
    # Page 2
    doc.add_heading('3. Flue Gas Emission Limits', level=2)
    tbl_emiss = doc.add_table(rows=3, cols=3)
    tbl_emiss.style = 'Table Grid'
    headers_emiss = ['Emission Parameter', 'Natural Gas Limit', 'Light Oil Limit']
    for i, h in enumerate(headers_emiss):
        tbl_emiss.cell(0, i).text = h
    emiss_rows = [
        ['CO Concentration', '< 50 mg/Nm³', '< 80 mg/Nm³'],
        ['NOx Class III', '< 100 mg/kWh', '< 150 mg/kWh'],
    ]
    for row_idx, r in enumerate(emiss_rows, start=1):
        for col_idx, val in enumerate(r):
            tbl_emiss.cell(row_idx, col_idx).text = val
            
    doc.save(str(docx_path))
    print(f'Generated: {docx_path}')

def generate_thermal_oil_txt():
    txt_path = DOCS_DIR / 'Thermal_Oil_Heater_Operating_Limits.txt'
    content = '''SELNIKEL THERMAL OIL HEATER (HOT OIL) TECHNICAL SPECIFICATION
Document: TOH-SPEC-2026 | Revision: REV-03 | Dept: Thermal Systems

1. Operating Fluid & Temperature Limits
Thermal fluid circuit must maintain continuous forced circulation across all firing rates.

Table: Thermal Oil Operational Limits
| Parameter | Minimum Value | Nominal Value | Maximum Safety Limit | Unit |
| --- | --- | --- | --- | --- |
| Flow Temperature | 180 °C | 280 °C | 320 °C | °C |
| System Operating Pressure | 2.5 bar | 4.0 bar | 6.0 bar | bar |
| Minimum Circulating Flow | 25 m³/h | 45 m³/h | 80 m³/h | m³/h |

2. Safety Interlocks
If fluid temperature exceeds 320 °C or pressure drops below 2.5 bar, emergency burner shutdown triggers automatically within 1.0 second.
'''
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(content.strip() + '\n')
    print(f'Generated: {txt_path}')

def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, 'rb') as f:
        while chunk := f.read(65536):
            h.update(chunk)
    return h.hexdigest()

def build_manifest():
    manifest = {
        "manifest_version": "1.2.0",
        "description": "Selnikel AI Synthetically Generated Industrial Technical Document Fixtures Inventory for Stage P1.2",
        "fixture_kind": "synthetic_generated",
        "synthetic": True,
        "review_status": "unverified_draft",
        "generator": "backend/scripts/generate_fixtures.py",
        "fixtures": [
            {
                "filename": "SB_Series_Steam_Boiler_Datasheet_REV01.pdf",
                "relative_path": "tests/fixtures/documents/SB_Series_Steam_Boiler_Datasheet_REV01.pdf",
                "format": "pdf",
                "sha256": sha256_file(DOCS_DIR / "SB_Series_Steam_Boiler_Datasheet_REV01.pdf"),
                "page_count": 3,
                "revision_code": "REV-01",
                "ocr_applied": False,
                "section_inventory": [
                    {"page_number": 1, "header": "1. General Overview & Design Standards"},
                    {"page_number": 2, "header": "2. Technical Operating Parameters & Capacities"},
                    {"page_number": 3, "header": "3. Safety Relief Valves & Auxiliary Limits"}
                ],
                "table_inventory": [
                    {
                        "table_id": "pdf_tab_01",
                        "page_number": 2,
                        "title": "Technical Operating Parameters & Capacities",
                        "column_count": 6,
                        "row_count": 4,
                        "headers": ["Model Code", "Steam Cap. (t/h)", "Working Press. (bar)", "Design Press. (bar)", "Steam Temp. (°C)", "Thermal Power (kW)"],
                        "ground_truth_cells": {
                            "SB-500": {"steam_cap": "0.45 t/h", "working_press": "10.0 bar", "design_press": "14.0 bar", "steam_temp": "184.1 °C", "thermal_power": "315 kW"},
                            "SB-1000": {"steam_cap": "0.9 t/h", "working_press": "14.0 bar", "design_press": "16.0 bar", "steam_temp": "198.3 °C", "thermal_power": "630 kW"},
                            "SB-2000": {"steam_cap": "1.8 t/h", "working_press": "14.0 bar", "design_press": "18.0 bar", "steam_temp": "198.3 °C", "thermal_power": "1260 kW"},
                            "SB-5000": {"steam_cap": "4.5 t/h", "working_press": "14.0 bar", "design_press": "22.0 bar", "steam_temp": "198.3 °C", "thermal_power": "3150 kW"}
                        }
                    },
                    {
                        "table_id": "pdf_tab_02",
                        "page_number": 3,
                        "title": "Safety Relief Valves & Auxiliary Limits",
                        "column_count": 4,
                        "row_count": 2,
                        "headers": ["Safety Device", "Set Pressure (bar)", "Connection DN (mm)", "Discharge Cap. (kg/h)"],
                        "ground_truth_cells": {
                            "Primary Safety Valve": {"set_pressure": "14.5 bar", "dn": "32 mm", "discharge_cap": "1100 kg/h"},
                            "Secondary Safety Valve": {"set_pressure": "15.0 bar", "dn": "40 mm", "discharge_cap": "1350 kg/h"}
                        }
                    }
                ]
            },
            {
                "filename": "SB_Series_Steam_Boiler_Datasheet.pdf",
                "relative_path": "tests/fixtures/documents/SB_Series_Steam_Boiler_Datasheet.pdf",
                "format": "pdf",
                "sha256": sha256_file(DOCS_DIR / "SB_Series_Steam_Boiler_Datasheet.pdf"),
                "page_count": 3,
                "revision_code": "REV-02",
                "ocr_applied": False,
                "section_inventory": [
                    {"page_number": 1, "header": "1. General Overview & Design Standards"},
                    {"page_number": 2, "header": "2. Technical Operating Parameters & Capacities"},
                    {"page_number": 3, "header": "3. Safety Relief Valves & Auxiliary Limits"}
                ],
                "table_inventory": [
                    {
                        "table_id": "pdf_tab_01",
                        "page_number": 2,
                        "title": "Technical Operating Parameters & Capacities",
                        "column_count": 6,
                        "row_count": 4,
                        "headers": ["Model Code", "Steam Cap. (t/h)", "Working Press. (bar)", "Design Press. (bar)", "Steam Temp. (°C)", "Thermal Power (kW)"],
                        "ground_truth_cells": {
                            "SB-500": {"steam_cap": "0.5 t/h", "working_press": "12.0 bar", "design_press": "16.0 bar", "steam_temp": "191.6 °C", "thermal_power": "350 kW"},
                            "SB-1000": {"steam_cap": "1.0 t/h", "working_press": "16.0 bar", "design_press": "18.0 bar", "steam_temp": "204.3 °C", "thermal_power": "700 kW"},
                            "SB-2000": {"steam_cap": "2.0 t/h", "working_press": "16.0 bar", "design_press": "20.0 bar", "steam_temp": "204.3 °C", "thermal_power": "1400 kW"},
                            "SB-5000": {"steam_cap": "5.0 t/h", "working_press": "16.0 bar", "design_press": "25.0 bar", "steam_temp": "204.3 °C", "thermal_power": "3500 kW"}
                        }
                    },
                    {
                        "table_id": "pdf_tab_02",
                        "page_number": 3,
                        "title": "Safety Relief Valves & Auxiliary Limits",
                        "column_count": 4,
                        "row_count": 2,
                        "headers": ["Safety Device", "Set Pressure (bar)", "Connection DN (mm)", "Discharge Cap. (kg/h)"],
                        "ground_truth_cells": {
                            "Primary Safety Valve": {"set_pressure": "16.5 bar", "dn": "32 mm", "discharge_cap": "1250 kg/h"},
                            "Secondary Safety Valve": {"set_pressure": "17.0 bar", "dn": "40 mm", "discharge_cap": "1500 kg/h"}
                        }
                    }
                ]
            },
            {
                "filename": "Monoblock_Burner_Maintenance_Manual.docx",
                "relative_path": "tests/fixtures/documents/Monoblock_Burner_Maintenance_Manual.docx",
                "format": "docx",
                "sha256": sha256_file(DOCS_DIR / "Monoblock_Burner_Maintenance_Manual.docx"),
                "page_count": 2,
                "revision_code": "REV-01",
                "ocr_applied": False,
                "section_inventory": [
                    {"page_number": 1, "header": "1. Technical Air & Combustion Data"},
                    {"page_number": 1, "header": "2. Periodic Maintenance Schedule & Intervals"},
                    {"page_number": 2, "header": "3. Diagnostic Fault Codes & Troubleshooting Matrix"},
                    {"page_number": 2, "header": "4. Burner Nozzle & Mechanical Parts Compatibility"}
                ],
                "table_inventory": [
                    {
                        "table_id": "docx_tab_01",
                        "page_number": 1,
                        "title": "Periodic Maintenance Schedule & Intervals",
                        "column_count": 4,
                        "row_count": 4,
                        "headers": ["Component", "Inspection Action", "Service Interval", "Torque / Spec"],
                        "ground_truth_cells": {
                            "Burner Nozzles": {"action": "Clean and inspect spray pattern", "interval": "500 hour", "torque": "25 Nm"},
                            "Nozzle Overhaul": {"action": "Replace atomizing disc and filter", "interval": "2000 hour", "torque": "30 Nm"},
                            "Gas Valve Train": {"action": "Tightness and leak test", "interval": "6 month", "torque": "No bubble at 100 mbar"},
                            "Complete Blower Overhaul": {"action": "Bearing lubrication and dynamic balancing", "interval": "1 year", "torque": "ISO 1940 G2.5"}
                        }
                    },
                    {
                        "table_id": "docx_tab_02",
                        "page_number": 2,
                        "title": "Diagnostic Fault Codes & Troubleshooting Matrix",
                        "column_count": 4,
                        "row_count": 3,
                        "headers": ["Fault Code", "Fault Description", "Possible Root Cause", "Action Required"],
                        "ground_truth_cells": {
                            "E01": {"desc": "Flame Failure Lockout", "cause": "Ionization probe contamination or gas cutoff", "action": "Clean ionization probe and verify supply pressure"},
                            "E04": {"desc": "Air Pressure Switch Fault", "cause": "Insufficient combustion air flow or clogged fan", "action": "Check blower differential pressure and air switch tube"},
                            "E07": {"desc": "High Flue Gas Temperature Trip", "cause": "Heat exchanger soot buildup or excess firing rate", "action": "Shutdown burner and perform boiler tube cleaning"}
                        }
                    }
                ]
            },
            {
                "filename": "Industrial_Boiler_Commissioning_Guide.docx",
                "relative_path": "tests/fixtures/documents/Industrial_Boiler_Commissioning_Guide.docx",
                "format": "docx",
                "sha256": sha256_file(DOCS_DIR / "Industrial_Boiler_Commissioning_Guide.docx"),
                "page_count": 2,
                "revision_code": "REV-01",
                "ocr_applied": False,
                "section_inventory": [
                    {"page_number": 1, "header": "1. Pre-Commissioning Electrical Limits"},
                    {"page_number": 1, "header": "2. Standards, Directives & Regulatory Compliance"},
                    {"page_number": 2, "header": "3. Flue Gas Emission Limits"}
                ],
                "table_inventory": [
                    {
                        "table_id": "docx_tab_01",
                        "page_number": 1,
                        "title": "Pre-Commissioning Electrical Limits",
                        "column_count": 3,
                        "row_count": 2,
                        "headers": ["Terminal", "Signal Type", "Allowed Range"],
                        "ground_truth_cells": {
                            "L1-L2-L3": {"signal": "Main Power", "range": "380 - 420 V"},
                            "P1-P2": {"signal": "Pressure 4-20 mA", "range": "0.0 - 25.0 bar"}
                        }
                    },
                    {
                        "table_id": "docx_tab_02",
                        "page_number": 2,
                        "title": "Flue Gas Emission Limits",
                        "column_count": 3,
                        "row_count": 2,
                        "headers": ["Emission Parameter", "Natural Gas Limit", "Light Oil Limit"],
                        "ground_truth_cells": {
                            "CO Concentration": {"ng": "< 50 mg/Nm³", "oil": "< 80 mg/Nm³"},
                            "NOx Class III": {"ng": "< 100 mg/kWh", "oil": "< 150 mg/kWh"}
                        }
                    }
                ]
            },
            {
                "filename": "Thermal_Oil_Heater_Operating_Limits.txt",
                "relative_path": "tests/fixtures/documents/Thermal_Oil_Heater_Operating_Limits.txt",
                "format": "txt",
                "sha256": sha256_file(DOCS_DIR / "Thermal_Oil_Heater_Operating_Limits.txt"),
                "page_count": 1,
                "revision_code": "REV-03",
                "ocr_applied": False,
                "section_inventory": [
                    {"page_number": 1, "header": "1. Operating Fluid & Temperature Limits"},
                    {"page_number": 1, "header": "2. Safety Interlocks"}
                ],
                "table_inventory": [
                    {
                        "table_id": "txt_tab_01",
                        "page_number": 1,
                        "title": "Thermal Oil Operational Limits",
                        "column_count": 5,
                        "row_count": 3,
                        "headers": ["Parameter", "Minimum Value", "Nominal Value", "Maximum Safety Limit", "Unit"],
                        "ground_truth_cells": {
                            "Flow Temperature": {"min": "180 °C", "nom": "280 °C", "max": "320 °C"},
                            "System Operating Pressure": {"min": "2.5 bar", "nom": "4.0 bar", "max": "6.0 bar"},
                            "Minimum Circulating Flow": {"min": "25 m³/h", "nom": "45 m³/h", "max": "80 m³/h"}
                        }
                    }
                ]
            }
        ]
    }
    with open(MANIFEST_PATH, 'w', encoding='utf-8') as f:
        json.dump(manifest, f, indent=2)
    ver = manifest['manifest_version']
    print(f'Generated: {MANIFEST_PATH} (version {ver})')

if __name__ == '__main__':
    generate_pdf_rev01()
    generate_pdf_rev02()
    generate_burner_docx()
    generate_commissioning_docx()
    generate_thermal_oil_txt()
    build_manifest()
